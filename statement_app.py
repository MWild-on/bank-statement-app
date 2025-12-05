# statement_app.py — раздел "Создание выписки"

import io
import zipfile
from copy import deepcopy
from datetime import datetime, date

import pandas as pd
import streamlit as st

import os
import tempfile
from docx2pdf import convert


from docx import Document   # pip install python-docx
from docx.shared import Pt
from docx.oxml.ns import qn

from ui_common import apply_global_css, section_header


# ===== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ТЕКСТА =====

def set_cell_text(cell, text, font_name="Times New Roman", font_size=10):
    """
    Записать текст в ячейку таблицы, сохранив шрифт Times New Roman.
    Полностью очищает содержимое ячейки и создаёт новый параграф.
    """
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")

    run.font.name = font_name
    run.font.size = Pt(font_size)

    # чтобы Word не подменял шрифт на Aptos
    r = run._element
    if r.rPr is None:
        r_rPr = r.makeelement(qn("w:rPr"))
        r.insert(0, r_rPr)
    else:
        r_rPr = r.rPr
    if r_rPr.rFonts is None:
        r_rFonts = r.makeelement(qn("w:rFonts"))
        r_rPr.insert(0, r_rFonts)
    else:
        r_rFonts = r_rPr.rFonts
    r_rFonts.set(qn("w:eastAsia"), font_name)


# ===== ПРОЧИЕ ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =====

def _format_date(d) -> str:
    if pd.isna(d) or d is None:
        return ""
    if isinstance(d, pd.Timestamp):
        d = d.date()
    return d.strftime("%d.%m.%Y")


def _format_amount(x) -> str:
    if pd.isna(x) or x == 0:
        return ""
    s = f"{float(x):,.2f}"          # 12345.6 -> 12,345.60

    # используем НЕРАЗРЫВНЫЙ пробел между тысячами
    nbsp = "\u00A0"
    s = s.replace(",", "X").replace(".", ",").replace("X", nbsp)

    return s



def _load_caseid_df(xls_bytes: bytes) -> pd.DataFrame:
    return pd.read_excel(io.BytesIO(xls_bytes), sheet_name="CaseID")


def _load_payments_raw(xls_bytes: bytes) -> pd.DataFrame:
    # читаем с сохранением текстовых колонок без преобразования
    return pd.read_excel(
        io.BytesIO(xls_bytes),
        sheet_name="Payments",
        header=[0, 1],
        dtype = {
            ("ДАННЫЕ ПО ВЫПИСКЕ", "№ документа"): str,
            ("ДАННЫЕ ПО ВЫПИСКЕ", "ВО"): str,
        }

# pd.read_excel(..., dtype=dtype)

    )


def _flatten_payments(raw: pd.DataFrame) -> pd.DataFrame:
    df = raw.iloc[1:].copy()  # пропускаем вторую строку шапки
    df.columns = [
        "_".join([str(c) for c in col if str(c) != "nan"]).strip()
        for col in df.columns
    ]
    return df


def _prepare_payments(df_flat: pd.DataFrame) -> pd.DataFrame:
    df = df_flat.copy()
    date_col = "ДАННЫЕ ПО ВЫПИСКЕ_Дата проводки"
    reg_col = "ДАННЫЕ ПО ВЫПИСКЕ_Рег.номер"
    debit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Сумма по дебету"
    credit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Сумма по кредиту"

    df[reg_col] = pd.to_numeric(df[reg_col], errors="coerce").astype("Int64")
    df[debit_col] = pd.to_numeric(df[debit_col], errors="coerce").fillna(0)
    df[credit_col] = pd.to_numeric(df[credit_col], errors="coerce").fillna(0)
    df["СумаПлатежа"] = df[debit_col] + df[credit_col]

    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    return df


def _fill_template(
    template_bytes: bytes,
    header_data: dict,
    payments_case_df: pd.DataFrame,
) -> bytes:
    """
    Сформировать один docx по шаблону (R или W) для одного Рег.номера.
    Все вставляемые тексты — Times New Roman.
    Из шапки меняем только дату/время и период, статические поля (банк, счёт,
    компания, валюта) берём из самого шаблона.
    """
    doc = Document(io.BytesIO(template_bytes))



def _docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Конвертировать DOCX (байты) в PDF (байты) через docx2pdf.
    Временные файлы удаляются после конвертации.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "tmp.docx")
        pdf_path = os.path.join(tmpdir, "tmp.pdf")

        # пишем docx во временный файл
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # конвертируем во временный pdf
        convert(docx_path, pdf_path)

        # читаем pdf в память
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes



    
    # ---------- ВЕРХНЯЯ ТАБЛИЦА ----------
    top = doc.tables[0]

    stmt_date_str = header_data["stmt_date"].strftime("%d.%m.%Y")
    stmt_time_str = header_data["stmt_time"].strftime("%H:%M:%S")

    period_str = (
        f"за период с {header_data['period_from'].strftime('%d.%m.%Y')} "
        f"по {header_data['period_to'].strftime('%d.%m.%Y')}"
    )

    # row 0, col 0: дата формирования (только дата)
    set_cell_text(top.cell(0, 0), stmt_date_str)

    # row 2: "Дата формирования выписки ..."
    msg = f"Дата формирования выписки {stmt_date_str} в {stmt_time_str}"
    set_cell_text(top.cell(2, 0), msg)
    set_cell_text(top.cell(2, 1), msg)

    # Период: row 5, col 2 и 3
    set_cell_text(top.cell(5, 2), period_str)
    set_cell_text(top.cell(5, 3), period_str)

    # Банк, счёт, компания, валюта НЕ трогаем — они зашиты в каждом шаблоне

    # ---------- НИЖНЯЯ ТАБЛИЦА (ОПЕРАЦИИ) ----------
    tbl = doc.tables[1]
    tbl_el = tbl._tbl

    # row 0 – заголовки, row 1 – подзаголовки, row 2 – шаблон строки
    template_row_tr = tbl.rows[2]._tr

    # очищаем строки ниже шаблонной (оставляем 0,1,2)
    while len(tbl.rows) > 3:
        tbl_el.remove(tbl.rows[-1]._tr)

    date_col = "ДАННЫЕ ПО ВЫПИСКЕ_Дата проводки"
    acc_debit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Счет"
    acc_credit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Счет.1"
    debit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Сумма по дебету"
    credit_col = "ДАННЫЕ ПО ВЫПИСКЕ_Сумма по кредиту"
    docnum_col = "ДАННЫЕ ПО ВЫПИСКЕ_№ документа"
    vo_col = "ДАННЫЕ ПО ВЫПИСКЕ_ВО"
    bank_col = "ДАННЫЕ ПО ВЫПИСКЕ_Банк (БИК и наименование)"
    purpose_col = "ДАННЫЕ ПО ВЫПИСКЕ_Назначение платежа"

    have_rows = False

    for _, r in payments_case_df.iterrows():
        new_tr = deepcopy(template_row_tr)
        tbl_el.append(new_tr)
        new_row = tbl.rows[-1]
        cells = new_row.cells

        set_cell_text(cells[0], _format_date(r[date_col]))
        set_cell_text(cells[1], r.get(acc_debit_col, "") or "")
        set_cell_text(cells[2], r.get(acc_credit_col, "") or "")
        set_cell_text(cells[3], _format_amount(r.get(debit_col, 0)))
        set_cell_text(cells[4], _format_amount(r.get(credit_col, 0)))
        set_cell_text(cells[5], r.get(docnum_col, "") or "")
        set_cell_text(cells[6], r.get(vo_col, "") or "")
        set_cell_text(cells[7], r.get(bank_col, "") or "")
        set_cell_text(cells[8], r.get(purpose_col, "") or "")

        have_rows = True

    # если добавили реальные строки — удаляем шаблонную (индекс 2)
    if have_rows and len(tbl.rows) > 3:
        try:
            tbl_el.remove(tbl.rows[2]._tr)
        except ValueError:
            pass

    out_buf = io.BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)
    return out_buf.getvalue()


def _update_caseid_with_sums(case_df: pd.DataFrame, payments_df: pd.DataFrame) -> pd.DataFrame:
    reg_col_payments = "ДАННЫЕ ПО ВЫПИСКЕ_Рег.номер"
    sums = payments_df.groupby(reg_col_payments)["СумаПлатежа"].sum()
    result = case_df.copy()
    result["Сума платежей"] = result["Рег.номер"].map(sums).fillna(0)
    return result


def _build_result_excel(case_df: pd.DataFrame, payments_raw: pd.DataFrame) -> io.BytesIO:
    """
    Итоговый Excel:
    - CaseID — плоский, без индекса
    - Payments — как в исходном файле (MultiIndex), index не убираем
    """
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        case_df.to_excel(writer, sheet_name="CaseID", index=False)
        payments_raw.to_excel(writer, sheet_name="Payments")
    out.seek(0)
    return out


# ===== ОСНОВНАЯ ФУНКЦИЯ РАЗДЕЛА =====

from ui_common import apply_global_css  # убедись, что импорт есть вверху файла

def run():
    apply_global_css()

    section_header(
        "Создание выписки",
        "Загрузите Excel-файл с листами CaseID и Payments. "
        "Параметры шапки выписки задаются ниже."
    )

    uploaded_xlsx = st.file_uploader(
        "Загрузите Excel с листами CaseID и Payments",
        type=["xlsx"],
    )

    # ----- ИНИЦИАЛИЗАЦИЯ ЗНАЧЕНИЙ ОДИН РАЗ -----
    if "stmt_date" not in st.session_state:
        st.session_state["stmt_date"] = date.today()

    if "stmt_time" not in st.session_state:
        st.session_state["stmt_time"] = datetime.now().time()

    if "period_from" not in st.session_state:
        today = date.today()
        st.session_state["period_from"] = today.replace(month=1, day=1)

    if "period_to" not in st.session_state:
        st.session_state["period_to"] = date.today()

    with st.form("stmt_header_form"):
        st.subheader("Параметры шапки выписки")

        stmt_date = st.date_input(
            "Дата формирования выписки",
            key="stmt_date",
        )

        stmt_time = st.time_input(
            "Время формирования выписки",
            key="stmt_time",
        )

        period_from = st.date_input(
            "Период с",
            key="period_from",
        )

        period_to = st.date_input(
            "Период по",
            key="period_to",
        )

        submitted = st.form_submit_button("Сформировать выписки")

    if not submitted:
        return


    if uploaded_xlsx is None:
        st.error("Нужно загрузить Excel с листами CaseID и Payments.")
        return

    # читаем шаблоны из репозитория
    template_r_path = "Template R.docx"
    template_w_path = "Template W.docx"

    template_r_bytes = None
    template_w_bytes = None

    try:
        with open(template_r_path, "rb") as f:
            template_r_bytes = f.read()
    except FileNotFoundError:
        st.error(f"Не найден файл шаблона: {template_r_path}.")
        return

    try:
        with open(template_w_path, "rb") as f:
            template_w_bytes = f.read()
    except FileNotFoundError:
        # W может не использоваться — тогда просто предупредим при попытке
        template_w_bytes = None

    xls_bytes = uploaded_xlsx.read()

    # читаем и готовим данные
    case_df = _load_caseid_df(xls_bytes)
    payments_raw = _load_payments_raw(xls_bytes)
    payments_flat = _flatten_payments(payments_raw)
    payments = _prepare_payments(payments_flat)

    # динамика шапки (одинакова для R и W)
    header_data = {
        "stmt_date": stmt_date,
        "stmt_time": stmt_time,
        "period_from": period_from,
        "period_to": period_to,
    }

    reg_col_payments = "ДАННЫЕ ПО ВЫПИСКЕ_Рег.номер"

    # считаем Excel с суммами
    case_updated = _update_caseid_with_sums(case_df, payments)
    result_excel = _build_result_excel(case_updated, payments_raw)

    # создаём один ZIP, куда кладём и выписки, и Excel
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        warning_w_missing = False

        # выписки по каждому Рег.номеру
        for _, row in case_df.iterrows():
            reg_num = row["Рег.номер"]
            template_code = str(row["Шаблон"]).strip().upper()

            if pd.isna(reg_num):
                continue

            try:
                reg_int = int(reg_num)
            except Exception:
                continue

            case_payments = payments[reg_col_payments] == reg_int
            case_payments = payments[case_payments]
            if case_payments.empty:
                continue

            if template_code == "R":
                template_bytes = template_r_bytes
            elif template_code == "W":
                if template_w_bytes is None:
                    warning_w_missing = True
                    continue
                template_bytes = template_w_bytes
            else:
                # неизвестный код шаблона — пропускаем
                continue

            docx_bytes = _fill_template(template_bytes, header_data, case_payments)
            
            # конвертация DOCX -> PDF
            pdf_bytes = _docx_bytes_to_pdf_bytes(docx_bytes)
            
            filename = f"{reg_int}_{template_code}.pdf"
            zf.writestr(filename, pdf_bytes)


        # добавляем Excel внутрь того же архива
        zf.writestr("CaseID_with_sums.xlsx", result_excel.getvalue())

    zip_buf.seek(0)

    st.success("Выписки сформированы.")

    st.download_button(
        "⬇️ Скачать архив (выписки + Excel)",
        data=zip_buf.getvalue(),
        file_name="statements_RW_with_excel.zip",
        mime="application/zip",
    )

    if template_w_bytes is None:
        st.warning("Шаблон W не найден в репозитории (Template W.docx). Выписки по W не сформированы.")
